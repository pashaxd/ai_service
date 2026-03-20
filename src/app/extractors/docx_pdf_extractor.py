from __future__ import annotations

from pathlib import Path
from typing import Any

from app.extractors.base import DocumentExtractor, ExtractedDocument


class DocxPdfExtractor(DocumentExtractor):
    """
    Extracts text from DOCX/PDF without Apache Tika (no Java / no jar downloads).

    - DOCX: paragraphs + tables (rows/cells).
    - PDF: per-page `extract_text()` via pdfplumber.
    """

    async def extract_text(self, file_path: str) -> ExtractedDocument:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            return self._extract_docx(file_path)
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        raise ValueError(f"Unsupported file type: {ext}")

    def _extract_docx(self, file_path: str) -> ExtractedDocument:
        from docx import Document  # type: ignore

        doc = Document(file_path)
        parts: list[str] = []

        # Paragraphs
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                parts.append(txt)

        # Tables (best-effort)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                row_txt = " | ".join([c for c in cells if c])
                if row_txt:
                    parts.append(row_txt)

        return ExtractedDocument(text="\n".join(parts).strip())

    def _extract_pdf(self, file_path: str) -> ExtractedDocument:
        import pdfplumber  # type: ignore

        def _pdf_text_to_diff_lines(extracted: str) -> list[str]:
            """
            Converts pdfplumber extracted text to diff-friendly lines.

            Key idea: DO NOT collapse the whole page into one paragraph (that kills
            diff granularity). Keep line breaks, and fix only hyphenated splits.
            """
            import re

            raw_lines = (extracted or "").splitlines()

            # First, clean and merge only hyphenated splits.
            cleaned: list[str] = []
            i = 0
            while i < len(raw_lines):
                ln = (raw_lines[i] or "").strip()
                if not ln:
                    i += 1
                    continue

                if ln.endswith("-") and i + 1 < len(raw_lines):
                    nxt = (raw_lines[i + 1] or "").strip()
                    if nxt:
                        ln = ln[:-1] + nxt
                        i += 2
                        cleaned.append(ln)
                        continue

                cleaned.append(ln)
                i += 1

            # Then, group into article-like blocks to reduce noise.
            # Most of your UI examples are based on "Статья <n>".
            re_article = re.compile(r"^\s*Статья\s+\d+(\.\d+)?\b", flags=re.IGNORECASE)
            blocks: list[str] = []
            cur_lines: list[str] = []
            cur_started = False

            for ln in cleaned:
                if re_article.match(ln):
                    if cur_started and cur_lines:
                        blocks.append("\n".join(cur_lines).strip())
                    cur_lines = [ln]
                    cur_started = True
                    continue
                if not cur_started:
                    # Skip headers/footers before first article.
                    continue
                # Preserve internal line breaks for better diff granularity.
                cur_lines.append(ln)

            if cur_started and cur_lines:
                blocks.append("\n".join(cur_lines).strip())

            # Fallback if no "Статья ..." found.
            return blocks if blocks else cleaned

        pages_text: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                txt: str | None = page.extract_text(x_tolerance=2, y_tolerance=2)  # type: ignore[assignment]
                txt = (txt or "").strip()
                if not txt:
                    continue
                diff_lines = _pdf_text_to_diff_lines(txt)
                # Don't insert page markers as diff lines: they add noise.
                pages_text.append("\n".join(diff_lines))

        return ExtractedDocument(text="".join(pages_text).strip(), meta={"pages": str(len(pages_text))})

